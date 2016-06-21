import docx
import pandas as pd
from .excel import Excel
import types

def Word(filepath_or_buffer=None):
    """
    Extends docx.document.Document
    """
    # patch class
    docx.document.Document.extract_text = extract_text
    docx.document.Document.extract_tables = extract_tables
    docx.document.Document.tables_to_excel = tables_to_excel
    docx.document.Document.add_tables_from_excel = add_tables_from_excel
    docx.document.Document.add_picture = add_picture

    # make instance using patched class
    doc = docx.Document(filepath_or_buffer)

    return doc

def extract_text(self):
    return '\n'.join([p.text for p in self.paragraphs])

def extract_tables(self):
    table_list = []
    for table in self.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ''
                for p in cell.paragraphs:
                    cell_text += p.text
                # if debug:
                #     print(cell_text, end=' ')
                row_data.append(cell_text)
            # if debug:
            #     print()
            table_data.append(row_data)

        try:
            import pandas as pd
            table_frame = pd.DataFrame(table_data)
            table_list.append(table_frame)
        except ImportError:
            # TODO: pandas가 없는 경우 텍스트로 출력
            pass

    return table_list[0] if len(table_list) == 1 else table_list

def tables_to_excel(self, filepath_or_buffer):
    table_list = self.extract_tables()

    excel_file = pd.ExcelWriter(filepath_or_buffer)
    for i, table_frame in enumerate(table_list):
        table_frame.to_excel(excel_file, 'Sheet{}'.format(i+1))
    excel_file.save()

    return filepath_or_buffer

def add_tables_from_excel(self, excel_filepath, sheetname=None,
    table_style='TableGrid'):
    table_frames = Excel(excel_filepath, sheetname=sheetname)

    if not isinstance(table_frames, list):
        table_frames = [table_frames]

    for frame in table_frames:
        frame = frame.fillna('')
        nrows, ncols = len(frame), len(frame.columns)
        table = self.add_table(rows=nrows, cols=ncols, style='TableGrid')
        # table.style = 'LightShading-Accent1'
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = str(frame.iloc[i, j])

def add_picture(self, image_path_or_stream, width=None, height=None):
    if width: width = docx.shared.Inches(width)
    if height: height = docx.shared.Inches(height)

    run = self.add_paragraph().add_run()
    return run.add_picture(image_path_or_stream, width, height)
